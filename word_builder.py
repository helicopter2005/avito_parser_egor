import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_tnr_12(paragraph):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def build_word_with_screenshots(data_rows, output_path):

    doc = Document()

    p1 = doc.add_paragraph(
        "Объявления о продаже объектов-аналогов, "
        "используемых в рамках сравнительного подхода."
    )
    for run in p1.runs:
        run.font.bold = True
    set_tnr_12(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # пустая строка

    p2 = doc.add_paragraph("Используемые в расчетах аналоги.")
    for run in p2.runs:
        run.font.bold = True
    set_tnr_12(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # отступ перед аналогами

    analogs = [r for r in data_rows if r["is_analog"]]
    offers = [r for r in data_rows if not r["is_analog"]]

    analog_counter = 1

    for row in analogs:
        analog = row["data"]
        title = analog.get("title")
        address = analog.get("address").replace('\n', ' ')
        url = analog.get("url", "")

        # --- Заголовок ---
        p_title = doc.add_paragraph(f"Объект аналог № {analog_counter}")
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in p_title.runs:
            run.font.bold = True

        set_tnr_12(p_title)

        # --- Путь к скриншотам ---
        screenshots_dir = os.path.join(
            os.getcwd(),
            "Скриншоты",
            title + address
        )

        price_history_img = os.path.join(screenshots_dir, "история цены.png")
        description_img = os.path.join(screenshots_dir, 'описание.png')
        publish_date_img = os.path.join(screenshots_dir, "дата_публикации.png")

        # --- Скриншоты ---
        if os.path.exists(price_history_img):
            p_img1 = doc.add_paragraph()
            p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img1.add_run()
            run.add_picture(price_history_img, width=Pt(500))
            set_tnr_12(p_img1)

        if os.path.exists(description_img):
            p_img2 = doc.add_paragraph()
            p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img2.add_run()
            run.add_picture(description_img, width=Pt(500))
            set_tnr_12(p_img2)

        if os.path.exists(publish_date_img):
            p_img3 = doc.add_paragraph()
            p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img3.add_run()
            run.add_picture(publish_date_img, width=Pt(500))
            set_tnr_12(p_img3)

        # --- Ссылка ---
        p_link = doc.add_paragraph(url)
        p_link.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_tnr_12(p_link)

        # Пустая строка между аналогами
        doc.add_paragraph("")

        analog_counter += 1

    if offers:
        p_offers = doc.add_paragraph("Предложения")
        p_offers.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_offers.runs:
            run.font.bold = True
        set_tnr_12(p_offers)

        doc.add_paragraph("")

        for row in offers:
            data = row["data"]

            title = data.get("title")
            address = data.get('address').replace('\n', ' ')
            url = data.get("url", "")

            screenshots_dir = os.path.join(
                os.getcwd(),
                "Скриншоты",
                title + address
            )

            price_history_img = os.path.join(screenshots_dir, "история цены.png")
            description_img = os.path.join(screenshots_dir, 'описание.png')
            publish_date_img = os.path.join(screenshots_dir, "дата_публикации.png")

            if os.path.exists(price_history_img):
                p_img1 = doc.add_paragraph()
                p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img1.add_run()
                run.add_picture(price_history_img, width=Pt(500))
                set_tnr_12(p_img1)

            if os.path.exists(description_img):
                p_img2 = doc.add_paragraph()
                p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img2.add_run()
                run.add_picture(description_img, width=Pt(500))
                set_tnr_12(p_img2)

            if os.path.exists(publish_date_img):
                p_img2 = doc.add_paragraph()
                p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img2.add_run()
                run.add_picture(publish_date_img, width=Pt(500))
                set_tnr_12(p_img2)

            p_link = doc.add_paragraph(url)
            set_tnr_12(p_link)

            doc.add_paragraph("")

    doc.save(output_path)
